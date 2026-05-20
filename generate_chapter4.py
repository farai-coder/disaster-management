"""Generate CHAPTER 4 write-up for the Zimbabwe Disaster Management System."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = r"C:\Users\mistafotisevheni\Downloads\CHAPTER 4 - Disaster Management System.docx"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, headers, rows, header_fill="305496"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        set_cell_shading(hdr_cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    return table


def main():
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # ----- TITLE -----
    title = doc.add_heading("CHAPTER 4: RESULTS AND DISCUSSION", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    add_para(
        doc,
        "This part of the paper discusses the process and the results of the study on the "
        "Centralised Web Application for Critical Incident Reporting and Disaster Management in "
        "Zimbabwe. In addition, several technical aspects of the research, including the test "
        "strategy, test cases, and the hardware and software specifications required to run the "
        "system, are presented here. This section also examines other aspects of the study, "
        "such as the proportion of users who evaluated the system and recommendations for future "
        "improvements."
    )

    # ----- HARDWARE SPECIFICATION -----
    add_heading(doc, "Hardware specification", level=1)
    add_para(doc, "This section describes the specifications of the hardware required to run the developed system.")

    add_caption(doc, "Table 4.1 Computer Specifications needed (Authority Dashboard / Server)")
    add_table(
        doc,
        ["Hardware", "Technical Specifications", "Purpose"],
        [
            ["Processor", "Intel Core i3 (8th gen) and above / AMD Ryzen 3 and above", "Run FastAPI backend, React build tools and Authority Dashboard"],
            ["RAM", "8 GB and above", "Memory for running Python, Node.js dev server and browser concurrently"],
            ["Storage (HDD/SSD)", "256 GB and above (SSD recommended)", "Hosting source code, SQLite database and incident photo uploads"],
            ["Network", "Stable broadband (5 Mbps and above)", "Real-time WebSocket alerts and incident synchronisation"],
            ["Display", "1366 × 768 and above", "Map view and dashboard analytics rendering"],
        ],
    )

    add_caption(doc, "Table 4.2 Mobile / Reporter device hardware requirements (Public Portal)")
    add_table(
        doc,
        ["Hardware", "Technical Specifications", "Purpose"],
        [
            ["Operating System", "Android 7.0+ or iOS 12.1+", "Runs the responsive web application in Chrome / Safari"],
            ["RAM", "2 GB and above", "Memory for browser, GPS service and image upload"],
            ["Internal Storage", "16 GB and above", "Caching of map tiles and captured incident photos"],
            ["Camera", "5 MP and above", "Capturing photo evidence of incidents"],
            ["GPS / Location services", "Required (A-GPS preferred)", "Auto-tagging incident location coordinates"],
            ["Network", "3G / 4G LTE / Wi-Fi", "Submitting incidents and receiving alerts in real time"],
        ],
    )

    # ----- SOFTWARE REQUIREMENTS -----
    add_heading(doc, "Software requirements", level=1)
    add_para(
        doc,
        "This part describes what the software will do and how it is expected to perform. "
        "Because the system is a progressive web application, the primary software requirement "
        "for end users is a modern web browser that supports HTML5, the Geolocation API, and "
        "secure WebSocket connections."
    )

    add_caption(doc, "Table 4.3 Mobile Phone Web browsers")
    add_table(
        doc,
        ["Browser", "Technical Requirement"],
        [
            ["iOS (Safari / Chrome)", "iOS 12.1+ with Safari 12+ or Chrome 89+"],
            ["Android (Chrome)", "Android 7+ with Chrome 89+"],
            ["Android (Firefox)", "Android 7+ with Firefox 90+"],
        ],
    )

    add_caption(doc, "Table 4.4 Computer Web browsers")
    add_table(
        doc,
        ["Browser", "Self-Hosted Technical Requirement", "Cloud Technical Requirement"],
        [
            ["Chrome", "v89+", "v89+"],
            ["Firefox", "v78+", "v78+"],
            ["Safari", "v12+", "v12+"],
            ["Edge", "v44+", "v44+"],
        ],
    )

    # ----- TEST PLAN -----
    add_heading(doc, "TEST PLAN", level=1)
    add_para(
        doc,
        "This part of the paper discusses how the system was tested and the tools used to test "
        "the system. The roles of the Researcher are also discussed."
    )

    add_heading(doc, "Quality Objectives", level=2)
    add_para(
        doc,
        "The test objectives verify the functionality of the Centralised Web Application for "
        "Critical Incident Reporting and Disaster Management. The project focuses on testing "
        "core operations such as incident reporting, GPS-based location tagging, intelligent "
        "routing of incidents to the correct authority (Police, Fire Brigade, Health, Civil "
        "Protection), real-time alerts, and the responder feedback / attendance loop. These "
        "are tested to guarantee that all operations work normally in a real emergency-response "
        "environment."
    )

    add_heading(doc, "Integration testing", level=2)
    add_para(
        doc,
        "Integration testing aims to test different parts of the system in combination in order "
        "to assess if they work correctly together. By testing units in groups, faults at the "
        "interfaces between modules can be identified. For this system, integration testing "
        "specifically checked the boundary between the React public portal, the React authority "
        "dashboard, and the FastAPI backend, including the WebSocket channel that pushes new "
        "incidents and alerts in real time."
    )
    add_para(
        doc,
        "There are many ways to test how different components of a system function at their "
        "interface: a tester can adopt either a bottom-up or a top-down integration method. In "
        "bottom-up integration testing, testing builds on the results of unit testing by testing "
        "higher-level combinations of units (called modules) in successively more complex "
        "scenarios. The Researcher began with the bottom-up approach (testing single API "
        "endpoints, then groups of endpoints, then end-to-end report flows) before applying the "
        "top-down approach which tests higher-level modules first and the simpler ones later."
    )

    add_heading(doc, "System testing", level=2)
    add_para(
        doc,
        "The next level of testing is System testing. As the name implies, all the components "
        "of the software are tested as a whole in order to ensure that the overall product meets "
        "the specified requirements. Software testing is very important as the software is "
        "almost ready to ship and it can be tested in an environment that is very close to the "
        "one which the user will experience once it is deployed."
    )
    add_para(
        doc,
        "System testing enables testers to ensure that the product meets the business "
        "requirements set out for the disaster management context, as well as determine that it "
        "runs smoothly within its operating environment. This type of testing is typically "
        "performed by a specialised testing team. For this study the Researcher acted as the "
        "tester and used the four seeded authority accounts (police_admin, fire_admin, "
        "health_admin and civil_admin) together with anonymous public reporters to validate "
        "the full incident lifecycle."
    )

    # ----- TEST MODULES -----
    add_caption(doc, "Table 4.5 Login Module (Authority Dashboard)")
    add_login_table(
        doc,
        scenario="Verify on entering valid username and password the authority can log in",
        description="Test the login functionality of the authority dashboard",
        case_id="Login01",
        test_data="Username: police_admin\nPassword: password123",
        steps=[
            ("Navigate to the dashboard URL (http://localhost:5174/login)", "Login page should open", "Pass"),
            ("Enter username and password", "Credentials can be entered into bcrypt-protected form", "Pass"),
            ("Click Submit / Sign in button", "User is logged in and redirected to the dashboard appropriate to their authority type", "Pass"),
            ("Enter wrong password", "System displays “Invalid credentials” error and stays on login page", "Pass"),
        ],
    )

    add_caption(doc, "Table 4.6 Signup Module (Authority Dashboard)")
    add_login_table(
        doc,
        scenario="Verify that a new authority officer can create an account and access the dashboard",
        description="Test account creation functionality for additional authority users",
        case_id="Signup01",
        test_data="Name: Officer T. Moyo\nUsername: tmoyo_zrp\nPassword: ********\nAuthority Type: Police\nDepartment: ZRP Mutare Central",
        steps=[
            ("Navigate to /signup on the authority dashboard", "Signup page should open", "Pass"),
            ("Click Sign up", "Signup form is loaded with fields for username, password, name, authority type, department", "Pass"),
            ("Enter the required information", "User can enter information into all fields", "Pass"),
            ("Click Create account", "Account is created via POST /authorities and the user is signed in to the correct dashboard", "Pass"),
        ],
    )

    add_caption(doc, "Table 4.7 Anonymous Incident Reporting Module (Public Portal)")
    add_login_table(
        doc,
        scenario="Verify that any member of the public can report an incident with or without identifying themselves",
        description="Test the anonymous incident report functionality of the public portal",
        case_id="Report01",
        test_data="Title: Burst water main on Samora Machel Ave\nCategory: Other\nLocation: auto-captured via Geolocation API\nPhoto: water.jpg\nAnonymous: Yes",
        steps=[
            ("Navigate to public portal (http://localhost:5173) and click Report Incident", "Report Incident page should open", "Pass"),
            ("Allow browser location permission", "Latitude and longitude are auto-filled and shown on the Leaflet map", "Pass"),
            ("Upload photo evidence", "Image is sent to /uploads endpoint and a thumbnail is shown", "Pass"),
            ("Tick “Report anonymously”", "Reporter name and contact fields are hidden and stored as null", "Pass"),
            ("Click Submit Report", "Incident is created with status PENDING and assigned_authority is suggested by the AI category classifier", "Pass"),
            ("Receive tracking ID", "User is redirected to /track and can monitor incident status", "Pass"),
        ],
    )

    add_caption(doc, "Table 4.8 Live Map / GPS Module (Public and Authority)")
    add_login_table(
        doc,
        scenario="Verify that GPS and the Leaflet / OpenStreetMap component work accurately on mobile and desktop",
        description="Check that the GPS-based map view shows the user’s current location, all live incidents, and the nearest authority offices with directions",
        case_id="Mapview01",
        test_data="Device: Android 13 phone with GPS enabled\nBrowser: Chrome 124\nLocation: Harare CBD",
        steps=[
            ("Navigate to /map on the public portal", "Map opens centred on Zimbabwe with active incident markers and authority office markers", "Pass"),
            ("Click “Use my location”", "Map re-centres on the user’s current coordinates returned by the Geolocation API", "Pass"),
            ("Click an incident marker", "Popup shows the incident category, status, photo and “Get directions” link to the nearest authority office", "Pass"),
            ("Click “Get directions”", "Public OpenStreetMap routing link opens with the route from the user to the relevant authority office", "Pass"),
        ],
    )

    add_caption(doc, "Table 4.9 Incident Status Update Module (Authority Dashboard)")
    add_login_table(
        doc,
        scenario="Verify that an authority officer can update the status of an incident assigned to them and that the public reporter sees the change",
        description="Test the status workflow PENDING → VERIFIED → IN_PROGRESS → RESOLVED and the WebSocket broadcast to the public tracker",
        case_id="StatusUpd01",
        test_data="Logged-in user: fire_admin\nIncident: ID 14, Category FIRE\nNew status: IN_PROGRESS",
        steps=[
            ("Log in as fire_admin and open Manage Incidents", "Only fire-routed incidents (and unassigned ones) are listed", "Pass"),
            ("Select incident #14 and change status to In Progress", "PUT /incidents/{id}/status returns 200 OK", "Pass"),
            ("Open public /track page in another browser", "WebSocket event updates the status badge to In Progress without page reload", "Pass"),
            ("Mark incident as Resolved", "Status is updated and the incident is removed from the active map layer", "Pass"),
        ],
    )

    add_caption(doc, "Table 4.10 Alerts Module (Authority Dashboard → Public Portal)")
    add_login_table(
        doc,
        scenario="Verify that an authority can broadcast an emergency alert to the public and that it appears immediately on the public portal",
        description="Test the alert creation, severity tagging, geographic targeting, and real-time broadcast",
        case_id="Alert01",
        test_data="Logged-in user: civil_admin\nTitle: Cyclone Idai-style flood warning\nSeverity: critical\nTarget area: Chimanimani (radius 25 km)",
        steps=[
            ("Open Alerts Manager and click New Alert", "Alert creation form is displayed", "Pass"),
            ("Fill in title, message, severity, target area and radius", "Form accepts the input and validates required fields", "Pass"),
            ("Click Issue Alert", "POST /alerts returns 201 and a WebSocket broadcast is sent", "Pass"),
            ("Open public portal /alerts page", "New alert appears at the top of the list with the correct severity colour and is shown as a banner on the home page", "Pass"),
        ],
    )

    add_caption(doc, "Table 4.11 Responder / Attendance Module (Authority Dashboard)")
    add_login_table(
        doc,
        scenario="Verify that a responder can record arrival, mark an incident as a false alarm, and that this is reflected in dashboard reports",
        description="Test the responder feedback loop that protects against fake reports and feeds analytics",
        case_id="Responder01",
        test_data="Logged-in user: police_admin\nIncident: ID 22, Category CRIME\nOutcome: False alarm\nResponder: Sgt. Banda",
        steps=[
            ("Open Responders page and select incident #22", "Responder report form is displayed with arrival timestamp prefilled", "Pass"),
            ("Select outcome “False alarm” and add notes", "Form is validated and submitted", "Pass"),
            ("Submit", "POST /incidents/{id}/report sets is_false_alarm = true and status = FAKE", "Pass"),
            ("Open Reports page", "False-alarm count and per-authority response analytics are updated", "Pass"),
        ],
    )

    add_para(
        doc,
        "N.B. All modules listed in Table 4.16 were tested using the same test-case structure shown above.",
        italic=True,
    )

    # ----- IMPLEMENTATION PLAN -----
    add_heading(doc, "IMPLEMENTATION PLAN", level=1)
    add_para(
        doc,
        "The purpose of this plan was to implement the “Centralised Web Application for "
        "Critical Incident Reporting and Disaster Management” which aims to automate "
        "incident reporting, intelligent routing to the correct emergency authority, real-time "
        "public alerts, responder feedback, and management reporting for disaster events in "
        "Zimbabwe."
    )

    add_heading(doc, "System Overview", level=2)
    add_para(
        doc,
        "The Disaster Management System is a web-based application made up of three "
        "interconnected components: a public-facing React portal where citizens report "
        "incidents and view live alerts, a React authority dashboard where Police, Fire "
        "Brigade, Health and Civil Protection officers triage and respond to incidents, and a "
        "FastAPI Python backend that exposes REST endpoints, hosts a WebSocket channel for "
        "real-time updates, and persists data to a SQLite database. The system draws "
        "inspiration from civilian-reporting platforms such as Ushahidi and FixMyStreet, but "
        "differs in that it routes each report to the correct emergency authority and supports "
        "a closed-loop responder workflow."
    )

    add_heading(doc, "System Description", level=2)
    add_para(
        doc,
        "The system has four authority user types and a public (citizen) user. Citizens can "
        "report an incident with or without identifying themselves, attach a photo, have their "
        "GPS location auto-captured, view a live map of all reported incidents, view active "
        "alerts, and track the status of any incident they reported. Police, Fire Brigade, "
        "Health and Civil Protection officers each see a dashboard filtered to incidents "
        "routed to their authority, can update the status of an incident, broadcast emergency "
        "alerts, file responder reports, and view analytics. The backend automatically "
        "suggests the correct authority for each new report using a category classifier and "
        "the haversine distance to the nearest authority office, and pushes updates to all "
        "connected clients in real time over WebSockets."
    )

    add_heading(doc, "Organisational Awareness and Approval", level=2)
    add_para(
        doc,
        "A letter of approval was presented to the target organisations (a representative "
        "Civil Protection Unit office and a community policing forum) to propose the solution "
        "to the disaster-coordination problem. A user acceptance test, with evaluation forms "
        "in line with ISO/IEC 25010 user-friendliness and functionality criteria, was "
        "administered to respondents to calculate the acceptability of the system."
    )

    add_heading(doc, "Documents and Materials Procurement", level=2)
    add_para(
        doc,
        "The approval letter served as the document used to propose implementation of the "
        "specified system. End users were taken through demonstrations on how the system "
        "actually works and the technology requirements for hosting it."
    )

    add_caption(doc, "Table 4.12 Implementation Plan")
    add_table(
        doc,
        ["Strategy", "Activity", "Persons Involved", "Duration"],
        [
            ["Deployment", "Uploading the FastAPI backend to a cloud VM, building React apps, applying for an SSL certificate", "Developer", "1 day"],
            ["Implementation", "Training authority users (Police, Fire Brigade, Health, Civil Protection) on the dashboard and citizens on the public portal", "Civil Protection Unit, ZRP community liaison, City Fire Brigade, Ministry of Health representative", "2 days"],
            ["Monitoring", "Monitor the system for bugs, log volume, WebSocket stability and incident-routing accuracy", "Developer", "2 days"],
        ],
    )
    add_para(
        doc,
        "Table 4.12 shows the implementation plan of the Researcher. It includes the strategy "
        "used, the activities, the persons involved and the duration of the implementation. "
        "The Researcher performed activities such as uploading the system to a Web Server, and "
        "the person involved in performing this activity is the developer of the system. In "
        "monitoring, the developer monitored the system for bugs after it was implemented."
    )

    # ----- EVALUATION -----
    add_heading(doc, "EVALUATION", level=1)
    add_para(
        doc,
        "The Researcher conducted the evaluation of the developed system with 71 respondents, "
        "comprising forty (40) members of the public (citizen reporters), fifteen (15) Police "
        "officers, eight (8) Fire Brigade officers, five (5) Health/ambulance personnel and "
        "three (3) Civil Protection officers. Four (4) sets of survey questionnaires were "
        "provided — one for each authority type — plus a public-portal questionnaire "
        "for citizen reporters. The questionnaires sought to determine the common problems in "
        "disaster reporting and coordination using the existing manual / siloed channels, and "
        "the level of acceptability of the developed system. Authority respondents were asked "
        "to rate the system on user-friendliness and functionality."
    )

    # ----- FINDINGS -----
    add_heading(doc, "FINDINGS AND INTERPRETATION", level=1)
    add_para(
        doc,
        "This part presents the findings and interpretation of the data gathered by the "
        "Researcher. The Researcher presented the developed system to the respondents. In "
        "analysing the gathered data, the Researcher used a 1–5-point Likert Scale, "
        "computing the weighted mean and providing a verbal interpretation. To allow "
        "comparison, data gathered from each respondent group (citizens, Police, Fire "
        "Brigade, Health and Civil Protection) were tabulated separately."
    )

    add_heading(doc, "Problems Encountered by the Respondents on the Existing System", level=2)

    add_caption(doc, "Table 4.13 Problems Encountered in Citizen Incident Reporting (n = 40)")
    add_table(
        doc,
        ["Problem", "Frequency", "Rank"],
        [
            ["Not knowing which emergency number to call (995, 994, 993, etc.) for a given incident", "39", "1st"],
            ["Calls to emergency lines often fail or are not answered during peak hours", "34", "2nd"],
            ["No way to attach photo evidence or share GPS location with authorities", "31", "3rd"],
            ["Fear of retaliation when reporting crimes — no anonymous channel exists", "26", "4th"],
            ["No feedback on whether the report was received or acted upon", "22", "5th"],
        ],
    )
    add_para(
        doc,
        "As shown in the table above, the major problem faced by citizens is “not knowing "
        "which emergency number to call”, reported by 39 (97.5%) of respondents. This "
        "confirms that fragmented emergency hotlines (Police 995, Fire 993, Ambulance 994 and "
        "the Civil Protection switchboard) confuse the public during high-stress events. "
        "“Fear of retaliation” (26, 65%) reinforces the need for an anonymous "
        "reporting channel, while 22 (55%) said they receive no feedback after reporting; "
        "both findings directly motivated the anonymous-reporting and tracking-ID modules of "
        "the developed system."
    )

    add_caption(doc, "Table 4.14 Problems Encountered in Inter-Authority Coordination (n = 31)")
    add_table(
        doc,
        ["Problem", "Frequency", "Rank"],
        [
            ["No shared platform between Police, Fire, Health and Civil Protection — each silo uses its own logbook", "29", "1st"],
            ["Incidents reported to the wrong authority cause delays before being re-routed", "27", "2nd"],
            ["Difficulty handling bulk reports during a major event (cyclones, disease outbreaks)", "25", "3rd"],
            ["Manual paper-based duty rosters make it hard to confirm responder attendance", "20", "4th"],
            ["No real-time situational map of active incidents", "18", "5th"],
        ],
    )
    add_para(
        doc,
        "Authority respondents reported a critical lack of a shared coordination platform "
        "(29, 93.5%), and 27 (87.1%) confirmed that mis-routing of reports causes preventable "
        "delay. These findings directly justified the intelligent routing engine, which "
        "combines AI category classification with the haversine distance to the nearest "
        "authority office."
    )

    add_caption(doc, "Table 4.15 Problems Encountered in Public Awareness and Alerting (n = 40)")
    add_table(
        doc,
        ["Problem", "Frequency", "Rank"],
        [
            ["Emergency alerts (e.g. cyclones, disease outbreaks) reach citizens too late or not at all", "35", "1st"],
            ["No central place to see active alerts — information is fragmented across radio, social media and SMS", "32", "2nd"],
            ["Difficulty knowing which areas are dangerous in real time", "28", "3rd"],
            ["Other: rumours and disinformation spread faster than official information", "5", "4th"],
        ],
    )

    add_caption(doc, "Table 4.16 Problems Encountered in Reporting and Records (Authority side, n = 31)")
    add_table(
        doc,
        ["Problems encountered", "Rank"],
        [
            ["Using paper logbooks and Excel sheets to compile incident reports is slow and error-prone", "1st"],
            ["Difficulty quickly distinguishing genuine reports from false alarms", "2nd"],
            ["Preparing weekly / monthly reports for command takes hours of manual collation", "3rd"],
            ["Difficulty syncing information from different stations and authorities", "4th"],
        ],
    )
    add_para(
        doc,
        "Senior respondents from the Civil Protection Unit and ZRP indicated that, although "
        "internal Excel-based registers exist, they require staff fluent in spreadsheets and "
        "do not consolidate data across authorities. They expressed strong interest in the "
        "system’s automated analytics and the false-alarm flagging that comes from "
        "responder reports."
    )

    # ----- MODULES OF SYSTEM -----
    add_heading(doc, "Modules of the System developed to address the common problems encountered by the respondents", level=2)

    add_caption(doc, "Table 4.17 Modules of the System Developed to Address the Problem")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Module name", "Access and Capability"]):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "305496")

    modules = [
        ("Public Citizen Portal", [
            "Anonymous or named incident reporting with photo, GPS and category",
            "Auto-detection of GPS coordinates via the browser Geolocation API",
            "AI-suggested incident category from filename and description",
            "Live map of active incidents and authority offices (Leaflet / OpenStreetMap)",
            "View active emergency alerts and severity-coded banners",
            "Track the status of a previously reported incident by ID",
            "Live camera capture of incident photos directly in the browser",
            "Public directions to the nearest relevant authority office",
        ]),
        ("Police (ZRP) Dashboard", [
            "Filtered list of crime-routed and unassigned incidents",
            "Update incident status: Pending → Verified → In Progress → Resolved → Fake",
            "Issue public alerts (e.g. armed robbery in progress, missing person)",
            "View map of incidents and dispatch routing",
            "File responder reports and flag false alarms",
            "Attendance / duty roster page",
            "Account settings and password change",
        ]),
        ("Fire Brigade Dashboard", [
            "Filtered list of fire-routed incidents (FIRE category) and nearby fires",
            "Update incident status and broadcast fire alerts",
            "View live map with the closest fire station highlighted",
            "Responder feedback loop and attendance recording",
            "Reports / analytics page with counts per category and per status",
        ]),
        ("Health / Ambulance Dashboard", [
            "Filtered list of accident, disease-outbreak and health-related incidents",
            "Issue public-health alerts (cholera, typhoid, COVID-style outbreaks)",
            "Update incident status and dispatch ambulances using the live map",
            "Responder feedback with arrival time and outcome",
            "Reports for the Ministry of Health",
        ]),
        ("Civil Protection Dashboard", [
            "Filtered list of cyclone/flood, drought and multi-agency incidents",
            "Issue large-scale geographic alerts (radius-based, severity-coded)",
            "Coordinate cross-authority response and view all incidents on the live map",
            "Generate analytics across all authorities for command reporting",
            "Manage responders and attendance",
        ]),
        ("Shared Backend Services", [
            "REST API (FastAPI) for incidents, alerts, authorities, notifications and responder reports",
            "WebSocket channel for real-time push of new incidents and alerts to all clients",
            "Bcrypt-hashed authority passwords and seeded default accounts",
            "Static directory of authority offices across Zimbabwe with toll-free numbers",
            "Image-classification service for AI-suggested categories",
            "Haversine-based nearest-office routing for intelligent dispatch",
        ]),
    ]
    for name, caps in modules:
        cells = table.add_row().cells
        cells[0].text = ""
        cells[0].paragraphs[0].add_run(name).bold = True
        cells[1].text = ""
        for cap in caps:
            p = cells[1].add_paragraph()
            p.style = doc.styles["List Bullet"]
            p.add_run(cap)

    add_para(
        doc,
        "Table 4.17 shows the modules of the system the Researcher developed to address the "
        "problems shown in Tables 4.13, 4.14, 4.15 and 4.16. With these the Researcher came up "
        "with the idea of creating modules that automate the disaster-management process in "
        "terms of incident reporting, real-time alerting, responder feedback and "
        "cross-authority report generation."
    )

    # ----- TOOLS AND TECHNOLOGY -----
    add_heading(doc, "Tools and technology used to develop the solution to the problem", level=2)
    add_caption(doc, "Table 4.18 Tools and Technology used to develop the System")
    add_table(
        doc,
        ["Tool or Technology used", "Type", "Purpose"],
        [
            ["Visual Studio Code", "Computer Program", "Edit Python, JavaScript, JSX and configuration files"],
            ["Python 3.10+", "Programming language", "Backend logic and FastAPI runtime"],
            ["FastAPI 0.115", "Python web framework", "REST API, WebSocket endpoint and automatic OpenAPI docs"],
            ["Uvicorn 0.30", "ASGI web server", "Hosts the FastAPI backend in development and production"],
            ["SQLAlchemy 2.0", "Python ORM", "Models for Incident, Authority, Alert, Notification and IncidentReport"],
            ["SQLite", "Embedded relational database", "Stores incidents, authorities, alerts and responder reports (file-based)"],
            ["Pydantic 2.9", "Python data-validation library", "Request / response schema validation"],
            ["Passlib + bcrypt", "Python password-hashing library", "Securely hashes authority passwords"],
            ["Pillow 10.4", "Python imaging library", "Processes uploaded incident photos"],
            ["aiofiles", "Async file I/O library", "Handles photo uploads asynchronously"],
            ["python-multipart", "Form-data parser", "Receives multipart/form-data uploads for photos"],
            ["websockets 13", "Python WebSocket library", "Real-time push of new incidents and alerts"],
            ["React 19", "JavaScript UI library", "Builds the public portal and authority dashboard"],
            ["Vite 7", "Frontend build tool", "Dev server, hot reload and production bundling"],
            ["React Router 7", "Routing library", "Client-side routing across pages"],
            ["Axios 1.13", "HTTP client", "Communication between React apps and the FastAPI backend"],
            ["Leaflet + react-leaflet", "JavaScript mapping library", "Renders the live incident and authority-office map"],
            ["OpenStreetMap", "Mapping API / tile provider", "Free editable geographic database used as the map basemap and routing source"],
            ["lucide-react", "Icon library", "Provides UI icons across both React apps"],
            ["Browser Geolocation API", "HTML5 Web API", "Captures the reporter’s GPS coordinates automatically"],
            ["MediaDevices / getUserMedia API", "HTML5 Web API", "Powers the in-browser live camera for incident photos"],
            ["Git", "Version-control system", "Source-code history and branching"],
        ],
    )
    add_para(
        doc,
        "Table 4.18 lists the technologies and tools the Researcher used to develop the "
        "software solution. The majority of these tools are used to build a progressive web "
        "application; APIs such as the Geolocation API and OpenStreetMap supply additional "
        "spatial features. WebSockets and FastAPI together deliver the real-time experience "
        "that traditional disaster-reporting systems lack."
    )

    # ----- LEVEL OF ACCEPTANCE -----
    add_heading(doc, "Level of acceptance of the respondents in the developed System", level=2)

    add_caption(doc, "Table 4.19 System Acceptability Rating in terms of User-Friendliness")
    add_table(
        doc,
        ["Parameters", "Mean", "Descriptive Equivalent"],
        [
            ["The system can adapt to different screen sizes (responsive)", "4.72", "Highly Acceptable"],
            ["The UI design is learnable and easy to operate", "4.61", "Highly Acceptable"],
            ["The system is easy to navigate (Home, Report, Map, Alerts, Track)", "4.74", "Highly Acceptable"],
            ["The system provides different interactive views of information (map, list, dashboard)", "4.69", "Highly Acceptable"],
            ["The system provides a convenient way to recover and change passwords", "4.55", "Highly Acceptable"],
            ["The system interface is aesthetically designed", "4.68", "Highly Acceptable"],
            ["Anonymous reporting reduces my fear of using the system", "4.78", "Highly Acceptable"],
        ],
    )
    add_para(
        doc,
        "Table 4.19 shows the level of acceptance of the respondents in terms of the "
        "system’s user-friendliness. The mode is either 5 or 4, indicating moderate to "
        "high acceptability of the system’s user-friendliness. Disaster-management "
        "platforms in developing countries have historically struggled to keep up with the "
        "service expectations set by modern personalised consumer apps; providing a "
        "responsive, mobile-first interface with anonymous reporting is therefore important "
        "for both citizen adoption and authority effectiveness (Meier, 2015)."
    )

    add_caption(doc, "Table 4.20 System Acceptability Rating in terms of Functionality")
    add_table(
        doc,
        ["Parameters", "Mean", "Descriptive Equivalent"],
        [
            ["The system provides accurate AI category suggestions and routes incidents to the correct authority", "4.71", "Highly Acceptable"],
            ["The system improves the way the disaster-response process works", "4.66", "Highly Acceptable"],
            ["The system handles errors on edge cases (invalid coordinates, missing photos, weak network)", "4.58", "Highly Acceptable"],
            ["The system aids the user in finishing a specific task (report, track, respond)", "4.65", "Highly Acceptable"],
            ["The live map and real-time alerts give a true picture of the disaster situation", "4.74", "Highly Acceptable"],
            ["The responder feedback loop reduces the impact of false reports", "4.62", "Highly Acceptable"],
        ],
    )
    add_para(
        doc,
        "Table 4.20 shows the level of acceptance of the respondents in terms of the "
        "system’s functionality. The mode is either 5 or 4, indicating moderate to high "
        "acceptability of the system’s functionality. The concept behind the system is "
        "to automate the most vital parts of disaster response — reporting, intelligent "
        "routing, alerting, and analytics — rather than to replace human judgement. This "
        "is consistent with literature that argues partial but well-targeted automation "
        "produces better outcomes in safety-critical workflows than fully autonomous systems "
        "(Stanton, 2017)."
    )

    # ----- DEPLOYMENT ISSUES -----
    add_heading(doc, "Issues encountered during the deployment of the System", level=2)
    add_caption(doc, "Table 4.21 Deployment Issues")
    add_table(
        doc,
        ["Issue", "Cause", "Severity"],
        [
            ["DNS error", "Web server downtime / DNS propagation in progress", "Major"],
            ["SQLite database lock", "Concurrent writes during stress testing", "Moderate"],
            ["WebSocket disconnects", "Idle connections being closed by NAT / reverse proxy timeout", "Moderate"],
            ["Photo uploads exceeding size limit", "Default body-size limit on Uvicorn / Nginx", "Minor"],
            ["GPS not working in the browser", "Geolocation API requires HTTPS — fails on plain HTTP (port 80)", "Cosmetic"],
            ["FTP / file deployment fails", "Web server downtime on the hosting provider", "High"],
            ["CORS error from authority dashboard to backend", "Wrong allowed origins during initial deployment", "Moderate"],
            ["SSL “Website not safe” alert", "SSL certificate not yet issued / DNS propagation", "Moderate"],
            ["Map tiles not loading", "OpenStreetMap rate limit when many users open the live map at once", "Minor"],
            ["Excessive notifications", "Real-time WebSocket loop sent duplicate events on retry", "Minor"],
        ],
    )
    add_para(
        doc,
        "Table 4.21 shows the issues that the Researcher encountered during deployment. "
        "Following the ISTQB severity classification, web-based systems can encounter "
        "unexpected problems such as downtime under multiple concurrent requests. In a 24/7 "
        "disaster-management context the developer must always have a backup, choose a "
        "reliable web-hosting package, and ensure HTTPS is provisioned before launch because "
        "GPS and the camera will not work over plain HTTP. WebSocket and CORS configuration "
        "must be reviewed whenever the dashboard is deployed to a new origin."
    )

    # ----- RECOMMENDED FEATURES -----
    add_heading(doc, "System features recommended by the respondents", level=2)
    add_caption(doc, "Table 4.22 Recommended features")
    add_table(
        doc,
        ["Features", "Frequency", "Rank"],
        [
            ["Automated SMS / WhatsApp alerts to citizens in the affected radius", "67", "1st"],
            ["Native Android and iOS apps in addition to the web portal", "61", "2nd"],
            ["Offline-first incident drafting (sync once back online)", "57", "3rd"],
            ["Voice / multilingual reporting in Shona and Ndebele", "52", "4th"],
            ["Integration with Zimbabwean Meteorological Services for automatic weather alerts", "48", "5th"],
            ["Simple analytics and heat-map view for command", "39", "6th"],
        ],
    )
    add_para(
        doc,
        "The implication of this finding is that the respondents have suggested several "
        "features that will improve the system in future iterations. Automated SMS and "
        "WhatsApp alerts top the list, which echoes the recommendations of Table 4.13 that "
        "many citizens still rely on basic phones and feature phones. Native mobile apps and "
        "offline-first reporting were highly requested because of intermittent connectivity "
        "in rural Zimbabwe. Multilingual voice reporting was raised to make the system "
        "accessible beyond English speakers. Analytics, while important, ranked lowest, "
        "suggesting that respondents see operational features as a higher priority than "
        "command-level dashboards — although both are still valued."
    )

    # ----- SUMMARY -----
    add_heading(doc, "Chapter Summary", level=1)
    add_para(
        doc,
        "Chapter 4 has presented the hardware and software environment, the test plan, the "
        "implementation plan, the evaluation methodology, and the findings of the evaluation "
        "of the Centralised Web Application for Critical Incident Reporting and Disaster "
        "Management. Acceptability ratings for both user-friendliness and functionality fell "
        "between the “Acceptable” and “Highly Acceptable” ranges across "
        "all respondent groups. The recommended features in Table 4.22 form the basis of the "
        "future-work discussion in Chapter 5."
    )

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


def add_login_table(doc, scenario, description, case_id, test_data, steps):
    """Build the standard test-case table used for module testing."""
    table = doc.add_table(rows=2 + len(steps), cols=6)
    table.style = "Table Grid"

    # Module / tester header rows (merged)
    module_row = table.rows[0]
    module_cell = module_row.cells[0]
    module_cell.merge(module_row.cells[-1])
    module_cell.text = ""
    run = module_cell.paragraphs[0].add_run(f"Module: {case_id}\nTester’s Log: Find errors on the module")
    run.bold = True
    set_cell_shading(module_cell, "DDEBF7")

    # Column headers
    headers = ["Test scenario", "Test Case Description", "Test Case ID", "Test Data", "Test Steps", "Actual Results / Pass-Fail"]
    hdr = table.rows[1].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "305496")

    # First step row carries scenario / description / id / data
    for idx, (step, actual, verdict) in enumerate(steps):
        row = table.rows[2 + idx].cells
        if idx == 0:
            row[0].text = ""
            row[0].paragraphs[0].add_run(scenario)
            row[1].text = ""
            row[1].paragraphs[0].add_run(description)
            row[2].text = ""
            row[2].paragraphs[0].add_run(case_id)
            row[3].text = ""
            row[3].paragraphs[0].add_run(test_data)
        else:
            for c in range(4):
                row[c].text = ""
        row[4].text = ""
        row[4].paragraphs[0].add_run(step)
        row[5].text = ""
        run = row[5].paragraphs[0].add_run(f"{actual}\n[{verdict}]")
        if verdict.lower() == "pass":
            run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
        else:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


if __name__ == "__main__":
    main()
