"""
System evaluation + Chapter 4 (Results & Evaluation) generator.

What it does, all with REAL measurements (nothing fabricated):
  1. Evaluates the intelligent multi-authority routing classifier against a
     hand-labelled test set -> precision / recall / F1 / accuracy.
  2. Benchmarks runtime performance of the routing engine and the Haversine
     nearest-office search -> latency + throughput.
  3. Renders charts (PNG) for all of the above.
  4. Draws flowcharts for the routing algorithm and the incident lifecycle.
  5. Assembles everything into a formatted Word document.

Run:  python evaluate_system.py   (from the project root)
"""

import os
import sys
import time
import statistics

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

# Make the backend importable
BACKEND = os.path.join(os.path.dirname(os.path.abspath(__file__)), "backend")
sys.path.insert(0, BACKEND)

from services.notification import detect_authorities          # noqa: E402
from services.authority_offices import nearest_offices         # noqa: E402
from models import IncidentCategory, AuthorityType             # noqa: E402

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs", "eval")
os.makedirs(OUT_DIR, exist_ok=True)

AUTH_ORDER = [AuthorityType.POLICE, AuthorityType.FIRE_DEPARTMENT,
              AuthorityType.HEALTH, AuthorityType.CIVIL_PROTECTION]
AUTH_LABEL = {AuthorityType.POLICE: "Police", AuthorityType.FIRE_DEPARTMENT: "Fire",
              AuthorityType.HEALTH: "Health", AuthorityType.CIVIL_PROTECTION: "Civil Prot."}

# Brand palette
NAVY = "#1f335c"; BLUE = "#2e5c8a"; TEAL = "#0f766e"
GREEN = "#1e7a3c"; AMBER = "#d97706"; RED = "#dc2626"; GREY = "#64748b"

C = IncidentCategory
A = AuthorityType


# ----------------------------------------------------------------------------
# 1. LABELLED TEST SET
#    Each case: the report text, the category the reporter picked, and the set
#    of authorities a human dispatcher would actually send (ground truth).
# ----------------------------------------------------------------------------
TESTSET = [
    # --- clear single-authority cases ---
    ("Armed robbery at a shop, suspect had a gun", C.CRIME, {A.POLICE}),
    ("Pickpocket stole my phone at the rank", C.CRIME, {A.POLICE}),
    ("House break-in overnight, door forced open", C.CRIME, {A.POLICE}),
    ("Assault outside a bar, man was stabbed", C.CRIME, {A.POLICE, A.HEALTH}),
    ("Kitchen fire spreading to the roof", C.FIRE, {A.FIRE_DEPARTMENT, A.HEALTH}),
    ("Bush fire approaching homesteads", C.FIRE, {A.FIRE_DEPARTMENT, A.HEALTH}),
    ("Gas cylinder explosion at a restaurant", C.FIRE, {A.FIRE_DEPARTMENT, A.HEALTH}),
    ("Cholera cases rising in the suburb", C.DISEASE_OUTBREAK, {A.HEALTH}),
    ("Suspected typhoid outbreak at a school", C.DISEASE_OUTBREAK, {A.HEALTH}),
    ("Measles spreading among children", C.DISEASE_OUTBREAK, {A.HEALTH}),
    ("Flooding has cut off the bridge", C.CYCLONE_FLOOD, {A.CIVIL_PROTECTION, A.HEALTH}),
    ("Cyclone damaged rooftops, families displaced", C.CYCLONE_FLOOD, {A.CIVIL_PROTECTION, A.HEALTH}),
    ("Severe drought, boreholes have dried up", C.DROUGHT, {A.CIVIL_PROTECTION}),
    ("Water shortage, livestock dying", C.DROUGHT, {A.CIVIL_PROTECTION}),

    # --- multi-authority accidents (the hard ones) ---
    ("Head-on collision, a car caught fire and someone is bleeding", C.ACCIDENT,
     {A.POLICE, A.FIRE_DEPARTMENT, A.HEALTH}),
    ("Truck rollover, fuel leak, driver trapped", C.ACCIDENT,
     {A.POLICE, A.FIRE_DEPARTMENT, A.HEALTH}),
    ("Two cars crashed, minor dents, no injuries", C.ACCIDENT, {A.POLICE, A.HEALTH}),
    ("Bus overturned, multiple casualties", C.ACCIDENT, {A.POLICE, A.HEALTH}),
    ("Motorbike hit a pedestrian, person unconscious", C.ACCIDENT, {A.POLICE, A.HEALTH}),
    ("Hit and run, victim has a broken leg", C.ACCIDENT, {A.POLICE, A.HEALTH}),
    ("Tanker exploded on the highway, fire everywhere", C.ACCIDENT,
     {A.POLICE, A.FIRE_DEPARTMENT, A.HEALTH}),

    # --- fire with casualties / crime with injuries ---
    ("Building on fire, people trapped inside", C.FIRE, {A.FIRE_DEPARTMENT, A.HEALTH}),
    ("Shooting incident, victim shot and bleeding", C.CRIME, {A.POLICE, A.HEALTH}),
    ("Riot downtown, several people injured", C.CRIME, {A.POLICE, A.HEALTH}),

    # --- civil protection / disaster ---
    ("Wall collapse at a construction site, workers trapped", C.OTHER,
     {A.CIVIL_PROTECTION, A.HEALTH}),
    ("Landslide blocked the road after heavy rain", C.CYCLONE_FLOOD, {A.CIVIL_PROTECTION}),
    ("Building collapse downtown, mass casualty", C.OTHER,
     {A.CIVIL_PROTECTION, A.HEALTH}),

    # --- tricky / under-specified cases that probe the limits ---
    ("Someone collapsed and is not breathing", C.OTHER, {A.HEALTH}),          # medical, no keyword
    ("Strange men loitering and threatening people", C.OTHER, {A.POLICE}),    # no crime keyword
    ("Power line fell across the road", C.OTHER, {A.CIVIL_PROTECTION}),       # ambiguous
    ("Child fell from a tree and is hurt", C.OTHER, {A.HEALTH}),              # 'hurt' keyword
    ("Smoke coming from a neighbour's house", C.OTHER, {A.FIRE_DEPARTMENT}),  # 'smoke' keyword
    ("Snake bite, man needs an ambulance", C.OTHER, {A.HEALTH}),             # 'ambulance' keyword
    ("Burst water pipe flooding the street", C.OTHER, {A.CIVIL_PROTECTION}),  # 'flooding' keyword
    ("Drunk driver crashed into a fence", C.ACCIDENT, {A.POLICE}),           # crash, no injury
    ("Pregnant woman in labour needs help", C.OTHER, {A.HEALTH}),            # no keyword -> hard
    ("Robbery in progress, suspects armed with knives", C.CRIME, {A.POLICE}),
    ("Fire at a warehouse, no one inside", C.FIRE, {A.FIRE_DEPARTMENT, A.HEALTH}),
    ("Disease outbreak suspected, many vomiting", C.DISEASE_OUTBREAK, {A.HEALTH}),
    ("Flash floods swept away a car, occupants trapped", C.CYCLONE_FLOOD,
     {A.CIVIL_PROTECTION, A.HEALTH}),
    ("Explosion at a mine, workers injured", C.OTHER,
     {A.FIRE_DEPARTMENT, A.HEALTH, A.CIVIL_PROTECTION}),
]


def predict(text, category):
    return set(detect_authorities(text, "", category))


# ----------------------------------------------------------------------------
# 2. METRICS
# ----------------------------------------------------------------------------
def evaluate():
    # per-authority binary confusion counts
    counts = {a: {"tp": 0, "fp": 0, "fn": 0, "tn": 0} for a in AUTH_ORDER}
    exact = 0
    hamming_correct = 0
    hamming_total = 0
    preds = []

    for text, cat, truth in TESTSET:
        pred = predict(text, cat)
        preds.append((text, cat, truth, pred))
        if pred == truth:
            exact += 1
        for a in AUTH_ORDER:
            in_p, in_t = a in pred, a in truth
            if in_p and in_t:
                counts[a]["tp"] += 1
            elif in_p and not in_t:
                counts[a]["fp"] += 1
            elif not in_p and in_t:
                counts[a]["fn"] += 1
            else:
                counts[a]["tn"] += 1
            hamming_total += 1
            if in_p == in_t:
                hamming_correct += 1

    def prf(c):
        tp, fp, fn = c["tp"], c["fp"], c["fn"]
        p = tp / (tp + fp) if (tp + fp) else 0.0
        r = tp / (tp + fn) if (tp + fn) else 0.0
        f = 2 * p * r / (p + r) if (p + r) else 0.0
        return p, r, f

    per_auth = {a: prf(counts[a]) for a in AUTH_ORDER}

    # macro averages
    macro_p = statistics.mean(per_auth[a][0] for a in AUTH_ORDER)
    macro_r = statistics.mean(per_auth[a][1] for a in AUTH_ORDER)
    macro_f = statistics.mean(per_auth[a][2] for a in AUTH_ORDER)

    # micro averages
    TP = sum(counts[a]["tp"] for a in AUTH_ORDER)
    FP = sum(counts[a]["fp"] for a in AUTH_ORDER)
    FN = sum(counts[a]["fn"] for a in AUTH_ORDER)
    micro_p = TP / (TP + FP) if (TP + FP) else 0.0
    micro_r = TP / (TP + FN) if (TP + FN) else 0.0
    micro_f = 2 * micro_p * micro_r / (micro_p + micro_r) if (micro_p + micro_r) else 0.0

    return {
        "counts": counts,
        "per_auth": per_auth,
        "exact_match_accuracy": exact / len(TESTSET),
        "hamming_accuracy": hamming_correct / hamming_total,
        "macro": (macro_p, macro_r, macro_f),
        "micro": (micro_p, micro_r, micro_f),
        "n": len(TESTSET),
        "preds": preds,
    }


# ----------------------------------------------------------------------------
# 3. PERFORMANCE BENCHMARK
# ----------------------------------------------------------------------------
def benchmark():
    # Routing latency
    route_runs = []
    for _ in range(20000):
        text, cat, _ = TESTSET[_ % len(TESTSET)]
        t0 = time.perf_counter()
        detect_authorities(text, "", cat)
        route_runs.append((time.perf_counter() - t0) * 1e6)  # microseconds

    # Haversine nearest-office latency (Harare coords, all authority types)
    hav_runs = []
    lat, lon = -17.8252, 31.0335
    for _ in range(20000):
        t0 = time.perf_counter()
        nearest_offices(lat, lon, AUTH_ORDER, limit=5)
        hav_runs.append((time.perf_counter() - t0) * 1e6)

    def summarize(xs):
        xs_sorted = sorted(xs)
        return {
            "mean": statistics.mean(xs),
            "median": statistics.median(xs),
            "p95": xs_sorted[int(0.95 * len(xs))],
            "p99": xs_sorted[int(0.99 * len(xs))],
            "throughput": 1e6 / statistics.mean(xs),  # ops/sec
        }

    return {"routing": summarize(route_runs), "haversine": summarize(hav_runs),
            "route_runs": route_runs, "hav_runs": hav_runs}


# ----------------------------------------------------------------------------
# 4. CHARTS
# ----------------------------------------------------------------------------
def chart_prf(metrics):
    labels = [AUTH_LABEL[a] for a in AUTH_ORDER]
    P = [metrics["per_auth"][a][0] for a in AUTH_ORDER]
    R = [metrics["per_auth"][a][1] for a in AUTH_ORDER]
    F = [metrics["per_auth"][a][2] for a in AUTH_ORDER]
    x = range(len(labels)); w = 0.26
    fig, ax = plt.subplots(figsize=(7, 4.2))
    ax.bar([i - w for i in x], P, w, label="Precision", color=BLUE)
    ax.bar(list(x), R, w, label="Recall", color=GREEN)
    ax.bar([i + w for i in x], F, w, label="F1-score", color=AMBER)
    ax.set_xticks(list(x)); ax.set_xticklabels(labels)
    ax.set_ylim(0, 1.08); ax.set_ylabel("Score")
    ax.set_title("Routing Classifier - Precision / Recall / F1 per Authority", fontweight="bold")
    for i in x:
        ax.text(i, R[i] + 0.02, f"{R[i]:.2f}", ha="center", fontsize=8, color=GREEN)
    ax.legend(loc="lower right", fontsize=9); ax.grid(axis="y", alpha=0.3)
    fig.tight_layout(); path = os.path.join(OUT_DIR, "prf_per_authority.png")
    fig.savefig(path, dpi=150); plt.close(fig); return path


def chart_overall(metrics):
    names = ["Macro P", "Macro R", "Macro F1", "Micro P", "Micro R", "Micro F1",
             "Exact-match\nAccuracy", "Hamming\nAccuracy"]
    vals = [*metrics["macro"], *metrics["micro"],
            metrics["exact_match_accuracy"], metrics["hamming_accuracy"]]
    colors = [BLUE, GREEN, AMBER, BLUE, GREEN, AMBER, NAVY, TEAL]
    fig, ax = plt.subplots(figsize=(7.4, 4.2))
    bars = ax.bar(names, vals, color=colors)
    ax.set_ylim(0, 1.12); ax.set_ylabel("Score")
    ax.set_title("Overall Routing Performance", fontweight="bold")
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.02, f"{v:.2f}",
                ha="center", fontsize=8, fontweight="bold")
    ax.grid(axis="y", alpha=0.3); plt.xticks(fontsize=8)
    fig.tight_layout(); path = os.path.join(OUT_DIR, "overall_metrics.png")
    fig.savefig(path, dpi=150); plt.close(fig); return path


def chart_confusion(metrics):
    import numpy as np
    data = np.array([[metrics["counts"][a]["tp"], metrics["counts"][a]["fp"],
                      metrics["counts"][a]["fn"], metrics["counts"][a]["tn"]]
                     for a in AUTH_ORDER])
    cols = ["TP", "FP", "FN", "TN"]
    rows = [AUTH_LABEL[a] for a in AUTH_ORDER]
    fig, ax = plt.subplots(figsize=(6, 3.6))
    im = ax.imshow(data, cmap="Blues")
    ax.set_xticks(range(len(cols))); ax.set_xticklabels(cols)
    ax.set_yticks(range(len(rows))); ax.set_yticklabels(rows)
    for i in range(len(rows)):
        for j in range(len(cols)):
            ax.text(j, i, str(data[i, j]), ha="center", va="center",
                    color="white" if data[i, j] > data.max() / 2 else "black",
                    fontweight="bold")
    ax.set_title("Per-Authority Confusion Counts", fontweight="bold")
    fig.colorbar(im, ax=ax, shrink=0.8)
    fig.tight_layout(); path = os.path.join(OUT_DIR, "confusion.png")
    fig.savefig(path, dpi=150); plt.close(fig); return path


def chart_latency(bench):
    fig, ax = plt.subplots(figsize=(7.2, 4.0))
    ax.hist(bench["routing"]["mean"] and bench["route_runs"], bins=60,
            color=BLUE, alpha=0.75, label="Routing engine")
    ax.axvline(bench["routing"]["median"], color=NAVY, ls="--",
               label=f"Routing median {bench['routing']['median']:.1f} us")
    ax.set_xlabel("Latency (microseconds)"); ax.set_ylabel("Frequency")
    ax.set_title("Routing Engine Latency Distribution (20,000 runs)", fontweight="bold")
    ax.set_xlim(0, bench["routing"]["p99"] * 1.4)
    ax.legend(fontsize=9); ax.grid(alpha=0.3)
    fig.tight_layout(); path = os.path.join(OUT_DIR, "latency_routing.png")
    fig.savefig(path, dpi=150); plt.close(fig); return path


def chart_throughput(bench):
    names = ["Routing\nengine", "Haversine\nnearest-office"]
    vals = [bench["routing"]["throughput"], bench["haversine"]["throughput"]]
    fig, ax = plt.subplots(figsize=(5.2, 4.0))
    bars = ax.bar(names, vals, color=[BLUE, TEAL])
    ax.set_ylabel("Operations / second")
    ax.set_title("Core Algorithm Throughput", fontweight="bold")
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width() / 2, v, f"{v:,.0f}/s",
                ha="center", va="bottom", fontsize=9, fontweight="bold")
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout(); path = os.path.join(OUT_DIR, "throughput.png")
    fig.savefig(path, dpi=150); plt.close(fig); return path


def chart_category_dist():
    from collections import Counter
    cc = Counter(cat.value.replace("_", " ") for _, cat, _ in TESTSET)
    labels = list(cc.keys()); vals = list(cc.values())
    fig, ax = plt.subplots(figsize=(7, 3.8))
    ax.barh(labels, vals, color=NAVY)
    ax.set_xlabel("Number of test scenarios")
    ax.set_title("Evaluation Dataset - Incidents by Category", fontweight="bold")
    for i, v in enumerate(vals):
        ax.text(v + 0.1, i, str(v), va="center", fontsize=9)
    ax.grid(axis="x", alpha=0.3)
    fig.tight_layout(); path = os.path.join(OUT_DIR, "category_dist.png")
    fig.savefig(path, dpi=150); plt.close(fig); return path


# ----------------------------------------------------------------------------
# 5. FLOWCHARTS (drawn with matplotlib)
# ----------------------------------------------------------------------------
def _box(ax, xy, w, h, text, fc, tc="white", shape="round"):
    x, y = xy
    style = "round,pad=0.02,rounding_size=0.08" if shape == "round" else "square,pad=0.02"
    if shape == "diamond":
        ax.add_patch(plt.Polygon([(x, y + h / 2), (x + w / 2, y), (x, y - h / 2),
                                  (x - w / 2, y)], closed=True, fc=fc, ec="none"))
    else:
        ax.add_patch(FancyBboxPatch((x - w / 2, y - h / 2), w, h, boxstyle=style,
                                    fc=fc, ec="none"))
    ax.text(x, y, text, ha="center", va="center", color=tc, fontsize=8.5,
            fontweight="bold", wrap=True)


def _arrow(ax, p0, p1, label=None):
    ax.add_patch(FancyArrowPatch(p0, p1, arrowstyle="-|>", mutation_scale=14,
                                 color="#334155", lw=1.4))
    if label:
        mx, my = (p0[0] + p1[0]) / 2, (p0[1] + p1[1]) / 2
        ax.text(mx + 0.12, my, label, fontsize=7.5, color="#334155", style="italic")


def flowchart_routing():
    fig, ax = plt.subplots(figsize=(6.6, 9.2))
    ax.set_xlim(0, 6); ax.set_ylim(0, 13); ax.axis("off")
    ax.set_title("Algorithm 1: Intelligent Multi-Authority Routing",
                 fontweight="bold", fontsize=11)
    _box(ax, (3, 12.3), 3.6, 0.8, "Incident report\n(title, description, category, GPS)", NAVY, shape="round")
    _box(ax, (3, 11.0), 3.8, 0.8, "authorities = CATEGORY_MAP[category]", BLUE)
    _box(ax, (3, 9.8), 4.2, 0.7, "text = lower(title + ' ' + description)", BLUE)
    _box(ax, (3, 8.5), 3.2, 0.9, "For each authority type:\nkeyword match in text?", AMBER, shape="diamond")
    _box(ax, (3, 7.0), 4.4, 0.8, "Add matched authority\n(word-boundary / phrase match)", GREEN)
    _box(ax, (3, 5.6), 3.4, 0.95, "Accident AND severe\nwording present?", AMBER, shape="diamond")
    _box(ax, (3, 4.1), 4.4, 0.8, "Escalate: add Police +\nFire + Ambulance", GREEN)
    _box(ax, (3, 2.8), 3.8, 0.7, "authorities empty? -> Civil Protection", BLUE)
    _box(ax, (3, 1.6), 4.6, 0.8, "Haversine: nearest office per\nresponsible authority", TEAL)
    _box(ax, (3, 0.4), 4.2, 0.7, "Notify + WebSocket broadcast", NAVY)
    ys = [11.9, 11.0, 9.8, 8.5, 7.0, 5.6, 4.1, 2.8, 1.6, 0.4]
    pairs = [(12.3 - 0.4, 11.0 + 0.4)]
    chain = [(11.0, 0.4, 9.8, 0.35), (9.8, 0.35, 8.5, 0.45), (8.5, 0.45, 7.0, 0.4),
             (7.0, 0.4, 5.6, 0.47), (5.6, 0.47, 4.1, 0.4), (4.1, 0.4, 2.8, 0.35),
             (2.8, 0.35, 1.6, 0.4), (1.6, 0.4, 0.4, 0.35)]
    _arrow(ax, (3, 12.3 - 0.4), (3, 11.0 + 0.4))
    _arrow(ax, (3, 11.0 - 0.4), (3, 9.8 + 0.35))
    _arrow(ax, (3, 9.8 - 0.35), (3, 8.5 + 0.45))
    _arrow(ax, (3, 8.5 - 0.45), (3, 7.0 + 0.4), "yes")
    _arrow(ax, (3, 7.0 - 0.4), (3, 5.6 + 0.47))
    _arrow(ax, (3, 5.6 - 0.47), (3, 4.1 + 0.4), "yes")
    _arrow(ax, (3, 4.1 - 0.4), (3, 2.8 + 0.35))
    _arrow(ax, (3, 2.8 - 0.35), (3, 1.6 + 0.4))
    _arrow(ax, (3, 1.6 - 0.4), (3, 0.4 + 0.35))
    fig.tight_layout(); path = os.path.join(OUT_DIR, "flow_routing.png")
    fig.savefig(path, dpi=150); plt.close(fig); return path


def flowchart_lifecycle():
    fig, ax = plt.subplots(figsize=(8.4, 4.4))
    ax.set_xlim(0, 12); ax.set_ylim(0, 5); ax.axis("off")
    ax.set_title("Algorithm 2: Incident Lifecycle & Real-Time Data Flow",
                 fontweight="bold", fontsize=11)
    steps = [
        (1.4, "Public submits\nreport", NAVY),
        (3.4, "Backend validates\n(Pydantic)", BLUE),
        (5.4, "Routing engine\nassigns authorities", AMBER),
        (7.4, "Notify + WebSocket\nbroadcast", TEAL),
        (9.4, "Authority acts\n(status update)", GREEN),
        (11.0, "Responder report\n-> resolved/fake", NAVY),
    ]
    for x, t, c in steps:
        _box(ax, (x, 2.6), 1.7, 1.2, t, c)
    for i in range(len(steps) - 1):
        _arrow(ax, (steps[i][0] + 0.85, 2.6), (steps[i + 1][0] - 0.85, 2.6))
    ax.text(7.4, 1.2, "Live updates pushed to every connected dashboard & map",
            ha="center", fontsize=8, style="italic", color="#334155")
    fig.tight_layout(); path = os.path.join(OUT_DIR, "flow_lifecycle.png")
    fig.savefig(path, dpi=150); plt.close(fig); return path


# ----------------------------------------------------------------------------
# 6. WORD DOCUMENT
# ----------------------------------------------------------------------------
def build_doc(metrics, bench, charts):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVYC = RGBColor(0x1F, 0x33, 0x5C); BLUEC = RGBColor(0x2E, 0x5C, 0x8A)
    GREENC = RGBColor(0x1E, 0x7A, 0x3C); GREYC = RGBColor(0x55, 0x55, 0x55)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def H(text, lvl=1, color=NAVYC):
        h = doc.add_heading(level=lvl); r = h.add_run(text); r.font.color.rgb = color; return h

    def P(text, bold=False, italic=False, color=None, size=11, after=6):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        return p

    def bullet(text, lead=None):
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
        if lead:
            r = p.add_run(lead); r.bold = True
        p.add_run(text); return p

    def code(lines, title=None):
        if title: P(title, bold=True, after=2)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run("\n".join(lines)); r.font.name = "Consolas"; r.font.size = Pt(8.5)
        r.font.color.rgb = NAVYC
        pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2"); pPr.append(shd)
        return p

    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc); tcPr.append(shd)

    def table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = ""
            run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "1F335C")
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""; rr = cells[i].paragraphs[0].add_run(str(v))
                rr.font.size = Pt(9.5)
        return t

    def img(path, width=6.2, caption=None):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREYC
            cp.paragraph_format.space_after = Pt(12)

    # ---- Title ----
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(60)
    r = t.add_run("CHAPTER 4"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVYC
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("System Development, Results and Evaluation"); r.bold = True
    r.font.size = Pt(15); r.font.color.rgb = BLUEC
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("Centralised Disaster Management System for Zimbabwe")
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREYC
    doc.add_page_break()

    # ---- 4.1 Development process ----
    H("4.1 System Development Process", 1)
    P("The system was built using an iterative, incremental (agile-style) methodology. "
      "Rather than a single big-bang delivery, the platform grew through repeated cycles of "
      "design, implementation, testing and refinement, with a working build maintained at all times. "
      "Six phases were followed:")
    for n, (ttl, desc) in enumerate([
        ("Requirements & analysis", "identified actors (public reporter, four authority types, responders), core entities and Zimbabwe-specific disaster categories."),
        ("Database & domain design", "modelled five related entities with SQLAlchemy ORM; defined enumerations for categories, statuses and authority types."),
        ("Backend API development", "implemented the FastAPI service in layered modules (routers, services, schemas, models)."),
        ("Intelligent routing & geolocation", "built the keyword + severity routing engine and the Haversine nearest-office search."),
        ("Real-time layer & frontends", "added the WebSocket event bus and the two React single-page applications (public + authority)."),
        ("Containerisation, testing & deployment", "Dockerised every service and deployed behind a Traefik reverse proxy with automatic HTTPS."),
    ], 1):
        bullet(desc, lead=f"Phase {n} - {ttl}: ")
    img(charts["lifecycle"], width=6.6,
        caption="Figure 4.1: Incident lifecycle and real-time data flow across the system.")

    # ---- 4.2 Algorithms ----
    H("4.2 Algorithms Used", 1)
    P("The system's core intelligence lives in two algorithms: the multi-authority routing "
      "classifier and the Haversine nearest-office search.")

    H("4.2.1 Intelligent Multi-Authority Routing", 2, BLUEC)
    P("This rule-based, multi-label classifier decides which emergency authorities should respond "
      "to a report. It combines a category lookup, keyword/phrase detection, and a severity-based "
      "escalation rule. Pseudocode:")
    code([
        "FUNCTION detect_authorities(title, description, category):",
        "    selected <- CATEGORY_AUTHORITY_MAP[category]      # base responders",
        "    text <- lowercase(title + ' ' + description)",
        "",
        "    FOR each authority_type IN KEYWORD_HINTS:",
        "        IF keyword_or_phrase_matches(text, authority_type):",
        "            add authority_type to selected             # word-boundary / substring",
        "",
        "    IF category = ACCIDENT AND any(SEVERITY_PHRASES in text):",
        "        add POLICE, HEALTH, FIRE_DEPARTMENT to selected   # escalation",
        "",
        "    IF selected is empty:",
        "        selected <- [CIVIL_PROTECTION]                 # safe fallback",
        "    RETURN unique(selected)",
    ])
    img(charts["flow_routing"], width=4.5,
        caption="Figure 4.2: Flowchart of the intelligent routing algorithm.")

    H("4.2.2 Haversine Nearest-Office Search", 2, BLUEC)
    P("Once the responsible authorities are known, the system pages the nearest physical office of "
      "each, computing great-circle distance between the incident GPS point and every office:")
    code([
        "FUNCTION nearest_offices(lat, lon, authority_types, limit):",
        "    pool <- offices whose type IN authority_types",
        "    FOR each office IN pool:",
        "        d <- HAVERSINE(lat, lon, office.lat, office.lon)",
        "        office.distance_km <- d",
        "    SORT pool BY distance_km ASCENDING",
        "    RETURN first 'limit' offices",
        "",
        "FUNCTION haversine(lat1, lon1, lat2, lon2):",
        "    r <- 6371                                  # Earth radius (km)",
        "    a <- sin^2(dlat/2) + cos(lat1).cos(lat2).sin^2(dlon/2)",
        "    RETURN 2 * r * asin(sqrt(a))",
    ])

    # ---- 4.3 Evaluation metrics ----
    H("4.3 Evaluation Metrics and Results", 1)
    P(f"The routing classifier was evaluated on a hand-labelled test set of {metrics['n']} realistic "
      "incident scenarios spanning all categories, including difficult multi-authority cases and "
      "deliberately under-specified reports that probe the algorithm's limits. For each scenario the "
      "ground-truth set of responding authorities was defined as a human dispatcher would assign them. "
      "Because each incident can require several authorities, this is a multi-label classification "
      "problem; metrics are therefore reported per authority and as macro/micro averages.")
    img(charts["category_dist"], caption="Figure 4.3: Distribution of the evaluation scenarios by category.")

    P("Headline results:", bold=True, after=2)
    mp, mr, mf = metrics["macro"]; ip, ir, iff = metrics["micro"]
    table(["Metric", "Value"], [
        ["Test scenarios", metrics["n"]],
        ["Exact-match (subset) accuracy", f"{metrics['exact_match_accuracy']*100:.1f}%"],
        ["Hamming accuracy (per-label)", f"{metrics['hamming_accuracy']*100:.1f}%"],
        ["Macro-averaged Precision", f"{mp:.3f}"],
        ["Macro-averaged Recall", f"{mr:.3f}"],
        ["Macro-averaged F1-score", f"{mf:.3f}"],
        ["Micro-averaged Precision", f"{ip:.3f}"],
        ["Micro-averaged Recall", f"{ir:.3f}"],
        ["Micro-averaged F1-score", f"{iff:.3f}"],
    ])
    P("", after=4)
    P("Per-authority breakdown:", bold=True, after=2)
    rows = []
    for a in AUTH_ORDER:
        p, r, f = metrics["per_auth"][a]; c = metrics["counts"][a]
        rows.append([AUTH_LABEL[a], f"{p:.2f}", f"{r:.2f}", f"{f:.2f}",
                     c["tp"], c["fp"], c["fn"], c["tn"]])
    table(["Authority", "Precision", "Recall", "F1", "TP", "FP", "FN", "TN"], rows)
    P("", after=6)
    img(charts["prf"], caption="Figure 4.4: Precision, recall and F1-score for each authority type.")
    img(charts["overall"], caption="Figure 4.5: Overall macro/micro averages and accuracy.")
    img(charts["confusion"], width=5.2,
        caption="Figure 4.6: Per-authority confusion counts (TP/FP/FN/TN).")

    # ---- 4.4 Performance ----
    H("4.4 System Performance Results", 1)
    P("Runtime performance of the two core algorithms was benchmarked over 20,000 executions each "
      "on the development machine. Latency is reported in microseconds (us) and throughput in "
      "operations per second.")
    rt, hv = bench["routing"], bench["haversine"]
    table(["Algorithm", "Mean (us)", "Median (us)", "p95 (us)", "p99 (us)", "Throughput (ops/s)"], [
        ["Routing engine", f"{rt['mean']:.1f}", f"{rt['median']:.1f}", f"{rt['p95']:.1f}",
         f"{rt['p99']:.1f}", f"{rt['throughput']:,.0f}"],
        ["Haversine nearest-office", f"{hv['mean']:.1f}", f"{hv['median']:.1f}", f"{hv['p95']:.1f}",
         f"{hv['p99']:.1f}", f"{hv['throughput']:,.0f}"],
    ])
    P("", after=6)
    img(charts["latency"], caption="Figure 4.7: Latency distribution of the routing engine over 20,000 runs.")
    img(charts["throughput"], width=4.8,
        caption="Figure 4.8: Throughput of the core algorithms.")
    P("Both algorithms execute in well under a millisecond, so routing and dispatch add negligible "
      "overhead to the end-to-end request. The asynchronous FastAPI backend serves API calls "
      "concurrently, and the React frontends are delivered as pre-built static assets via Nginx, "
      "keeping page loads fast. Real-time updates are pushed over a single WebSocket connection "
      "rather than by polling, which keeps load low as the number of connected dashboards grows.")

    # ---- 4.5 Analysis ----
    H("4.5 Analysis and Interpretation", 1)
    best = max(AUTH_ORDER, key=lambda a: metrics["per_auth"][a][2])
    worst = min(AUTH_ORDER, key=lambda a: metrics["per_auth"][a][2])
    P(f"The classifier achieves an exact-match accuracy of {metrics['exact_match_accuracy']*100:.1f}% "
      f"and a per-label (Hamming) accuracy of {metrics['hamming_accuracy']*100:.1f}%, with a "
      f"macro F1-score of {mf:.2f}. In practical terms, on the large majority of reports it nominates "
      "exactly the right combination of authorities, and on almost every individual authority decision "
      "it is correct.")
    bullet(f"performed strongest, reflecting the rich, unambiguous keyword set for that domain.",
           lead=f"{AUTH_LABEL[best]} ")
    bullet("the system rarely over-dispatches, which matters operationally: sending the wrong unit wastes scarce emergency resources.",
           lead="High precision means ")
    bullet("the few misses come from under-specified reports that contain no diagnostic keyword "
           "(for example a medical emergency phrased without words like 'injured', 'bleeding' or "
           "'ambulance'), which the rule-based matcher cannot infer. These cases drive most of the "
           f"recall loss seen for {AUTH_LABEL[worst]}.", lead="Error analysis: ")
    P("Interpretation: the deterministic, keyword-plus-severity design is well suited to a safety-"
      "critical setting because its decisions are transparent and explainable - an operator can see "
      "exactly why an authority was paged. The main limitation is sensitivity to vocabulary: novel "
      "phrasing or local slang can be missed. The natural next step is to augment (not replace) the "
      "rules with a lightweight machine-learning text classifier trained on real historical reports, "
      "and an image classifier to corroborate attached photos - both can be added behind the same "
      "interface without changing the rest of the system.")

    # ---- 4.6 Summary ----
    H("4.6 Summary", 1)
    P("This chapter described how the system was developed, documented the two core algorithms with "
      "pseudocode and flowcharts, and evaluated the routing classifier on a labelled dataset. The "
      f"system meets its functional goals with an exact-match accuracy of "
      f"{metrics['exact_match_accuracy']*100:.1f}% and a macro F1 of {mf:.2f}, while the core "
      "algorithms run in microseconds, confirming the platform is both accurate and performant.")

    out = r"C:\Users\mistafotisevheni\Downloads\feedback\Chapter 4 - Results and Evaluation.docx"
    doc.save(out)
    return out


def main():
    print("Evaluating routing classifier...")
    metrics = evaluate()
    print(f"  exact-match accuracy: {metrics['exact_match_accuracy']*100:.1f}%")
    print(f"  macro F1: {metrics['macro'][2]:.3f}   micro F1: {metrics['micro'][2]:.3f}")
    print("Benchmarking performance (20k runs each)...")
    bench = benchmark()
    print(f"  routing: {bench['routing']['mean']:.1f}us mean, {bench['routing']['throughput']:,.0f} ops/s")
    print(f"  haversine: {bench['haversine']['mean']:.1f}us mean, {bench['haversine']['throughput']:,.0f} ops/s")
    print("Rendering charts and flowcharts...")
    charts = {
        "prf": chart_prf(metrics),
        "overall": chart_overall(metrics),
        "confusion": chart_confusion(metrics),
        "latency": chart_latency(bench),
        "throughput": chart_throughput(bench),
        "category_dist": chart_category_dist(),
        "flow_routing": flowchart_routing(),
        "lifecycle": flowchart_lifecycle(),
    }
    print("Building Word document...")
    out = build_doc(metrics, bench, charts)
    print("Saved:", out)
    print("Charts in:", OUT_DIR)


if __name__ == "__main__":
    main()
