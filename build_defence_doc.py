"""Generate a formatted Word version of the Project Defence Guide."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x33, 0x5C)
BLUE = RGBColor(0x2E, 0x5C, 0x8A)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, white=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color

def add_heading(text, level=1, color=NAVY):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    return h

def add_para(text, bold=False, italic=False, color=None, size=11, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_number(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows, header_fill="1F335C", widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_text(hdr[i], htext, bold=True, white=True, size=10)
        shade_cell(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

# ----------------------------------------------------------------------
# TITLE PAGE
# ----------------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(120)
r = title.add_run("PROJECT DEFENCE GUIDE")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Centralised Disaster Management System for Zimbabwe")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = BLUE

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.paragraph_format.space_before = Pt(24)
r = tag.add_run("A plain-language write-up of what the system is, how it was built,\n"
                "how it was deployed, the challenges faced, and how to defend it —\n"
                "organised to match the Final Assessment Form (June 2026).")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(160)
r = meta.add_run("Capstone Project  |  Final Assessment Defence Preparation")
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_page_break()

# ----------------------------------------------------------------------
# 1. PROBLEM
# ----------------------------------------------------------------------
add_heading("1. The Problem (Why this project exists)", 1)
add_para("In Zimbabwe, when a disaster or emergency happens — a fire, a road accident, a "
         "robbery, a disease outbreak, a flood — members of the public often do not know who "
         "to call, calls go to the wrong department, or several departments are needed at once "
         "but only one is contacted. There is no single national platform where:")
add_bullet("the public can report an incident with its exact location and a photo,")
add_bullet("the right authorities are alerted automatically and immediately,")
add_bullet("authorities can see all incidents on a live map, manage them, and issue public warnings,")
add_bullet("responders can report back what they found on the scene (genuine emergency vs. false alarm).")
add_para("My system solves this by providing one centralised web platform that connects the public "
         "and the emergency authorities (Police, Fire Brigade, Health/Ambulance, Civil Protection) in real time.")
add_para("One-sentence pitch for the panel:", bold=True, space_after=2)
add_para("“It is a centralised web application that lets any member of the public report a "
         "critical incident in seconds, automatically routes it to the correct emergency authorities "
         "based on what happened and where, and gives those authorities a live dashboard and map to "
         "respond, warn the public, and close the loop.”", italic=True, color=BLUE)

# ----------------------------------------------------------------------
# 2. WHAT IT DOES
# ----------------------------------------------------------------------
add_heading("2. What the System Does (Core Features)", 1)
add_para("The platform is made of three separate applications that talk to one shared backend:")

add_heading("A. Public Portal  (disaster.shopivell.com)", 2, BLUE)
add_bullet(" choose a category (crime, fire, accident, disease outbreak, cyclone/flood, drought, other), "
           "write a description, pin the location on a map (or use GPS), optionally attach a photo. Can be submitted anonymously.",
           "Report an Incident —")
add_bullet(" see incidents happening around the country in real time.", "Live Map —")
add_bullet(" follow the status of something you reported (pending → verified → in progress → resolved).", "Track an Incident —")
add_bullet(" read public safety warnings issued by the authorities.", "View Alerts —")
add_bullet(" Zimbabwe toll-free directory (Police 995, Ambulance 994, Fire 993, Childline 116, GBV/Musasa, ZESA, ZINWA).", "Emergency Numbers —")

add_heading("B. Authority Dashboard  (admin.disaster.shopivell.com)", 2, BLUE)
add_bullet(" for the four authority types.", "Login / Sign-up —")
add_bullet(" statistics scoped to that authority (total, pending, verified, in-progress, resolved, fake; active alerts; breakdown by category; recent incidents).", "Dashboard —")
add_bullet(" view, filter, update status, reassign, and delete incidents relevant to that department.", "Manage Incidents —")
add_bullet(" geographic view of all relevant incidents.", "Incident Map —")
add_bullet(" create, deactivate, delete public warnings (severity low/medium/high/critical, target area, radius, expiry).", "Alerts Manager —")
add_bullet(" track which responders attended which incident; validate and close reports.", "Responders & Attendance —")
add_bullet(" printable summaries aggregated by category, status and day, with date filtering.", "Reports / Analytics —")
add_bullet(" change password.", "Account Settings —")
add_bullet(" a bell that updates instantly when a new relevant incident arrives.", "Live notifications —")

add_heading("C. Backend API  (api.disaster.shopivell.com)", 2, BLUE)
add_para("The brain of the system — a REST API plus a WebSocket channel for live updates, with a database underneath.")

# ----------------------------------------------------------------------
# 3. HOW BUILT
# ----------------------------------------------------------------------
add_heading("3. How It Was Built (Step by Step)", 1)
add_para("This is the chronological story you can tell the panel.")
add_number(" Identified the actors (public reporter, four authority types, responders), the core entities "
           "(Incident, Authority, Alert, Notification, Incident Report), and the disaster categories relevant to Zimbabwe.",
           "Requirements & design. ")
add_number(" Designed the data model and implemented it with SQLAlchemy ORM so the schema is defined in Python code "
           "(models.py) and tables are created automatically.", "Database design. ")
add_number(" Built the server with FastAPI (Python), organised into routers/ (endpoints), services/ (business logic), "
           "schemas.py (Pydantic validation), and models.py / database.py (data layer).", "Backend API. ")
add_number(" The standout feature: instead of one fixed department, the system reads the title and description, matches "
           "keywords, combines that with the category to decide all responsible authorities, then finds the nearest "
           "physical office of each using the Haversine distance formula.", "Smart routing engine. ")
add_number(" Added a WebSocket endpoint so that when anything changes (new incident, status update, new alert) the server "
           "broadcasts the event and every connected dashboard and map updates instantly — no refresh needed.", "Real-time layer. ")
add_number(" Built the public portal and authority dashboard as two separate React 19 + Vite single-page apps, using React "
           "Router, Axios, Leaflet/React-Leaflet maps, and Lucide icons. A shared useLiveIncidents hook subscribes to the "
           "WebSocket and auto-reconnects with exponential backoff.", "Frontend apps. ")
add_number(" On first startup the backend auto-creates four demo authority accounts and runs lightweight automatic "
           "migrations to add new columns without losing data.", "Seeding & convenience. ")
add_number(" Wrote a Dockerfile for each app and a docker-compose.yml to run all three together behind a reverse proxy with HTTPS.",
           "Containerisation & deployment. ")

# ----------------------------------------------------------------------
# 4. DEPLOYMENT
# ----------------------------------------------------------------------
add_heading("4. How It Was Deployed", 1)
add_para("Local development first.", bold=True, space_after=2)
add_bullet(" uvicorn main:app --reload --port 8000  (API docs auto-generated at /docs).", "Backend:")
add_bullet(" npm run dev on port 5173.", "Public frontend:")
add_bullet(" npm run dev on port 5174.", "Authority dashboard:")
add_para("Production deployment with Docker.", bold=True, space_after=2)
add_para("Every component is packaged into a Docker container:")
add_bullet(" built from python:3.12-slim, installs dependencies, runs Uvicorn on port 8000.", "Backend image —")
add_bullet(" multi-stage build: a Node stage compiles the React app into static files, then a tiny Nginx image serves them — keeping the final image small and fast.", "Each frontend —")
add_para("docker-compose.yml orchestrates all three services behind a Traefik reverse proxy, which routes each "
         "domain to the right container, automatically obtains and renews HTTPS certificates from Let’s Encrypt, "
         "and force-redirects all HTTP traffic to HTTPS. The result is three live, secured domains:")
add_table(["Application", "Domain"],
          [["Public portal", "https://disaster.shopivell.com"],
           ["Authority dashboard", "https://admin.disaster.shopivell.com"],
           ["API", "https://api.disaster.shopivell.com"]],
          widths=[2.5, 3.5])
add_para("")
add_para("The SQLite database and uploaded photos are stored on Docker volumes, so data survives container "
         "restarts and re-deploys (this is also your basic backup/persistence answer).")
add_para("One-line deploy story for the panel:", bold=True, space_after=2)
add_para("“Each app is a Docker container; Docker Compose runs them together behind a Traefik reverse proxy "
         "that gives free auto-renewing HTTPS, and the database lives on a persistent volume so data is never lost on redeploy.”",
         italic=True, color=BLUE)

# ----------------------------------------------------------------------
# 5. ARCHITECTURE
# ----------------------------------------------------------------------
add_heading("5. System Architecture & Database Design", 1)
add_para("Architecture (three-tier)", bold=True, space_after=2)
code_block([
 "   PUBLIC USERS                      AUTHORITY STAFF",
 "        |                                  |",
 "        v                                  v",
 " +--------------+                 +------------------+",
 " | Public React |                 | Authority React  |   (Presentation tier)",
 " |   Portal     |                 |   Dashboard      |",
 " +------+-------+                 +--------+---------+",
 "        |   HTTPS (REST) + WebSocket (live)|",
 "        +---------------+------------------+",
 "                        v",
 "              +--------------------+",
 "              |  FastAPI Backend   |   (Application / logic tier)",
 "              |  routers + services|",
 "              |  smart routing,    |",
 "              |  WebSocket hub     |",
 "              +---------+----------+",
 "                        v",
 "              +--------------------+",
 "              |  SQLite Database   |   (Data tier)",
 "              |  via SQLAlchemy ORM|",
 "              +--------------------+",
])
add_para("Database tables (entities)", bold=True, space_after=4)
add_table(["Table", "Purpose", "Key fields"],
  [["incidents", "A reported emergency", "title, description, category, status, latitude, longitude, location_name, photo_url, is_anonymous, reporter info, assigned_authority, timestamps"],
   ["authorities", "An emergency-department account", "username, password_hash, name, authority_type, department, is_active"],
   ["alerts", "A public safety warning", "title, message, severity, category, target_area, lat/lon, radius_km, is_active, expires_at"],
   ["notifications", "In-app alert to a department", "message, type, target_authority_type, is_read, links to incident/alert"],
   ["incident_reports", "Responder feedback after attending", "responder_name, authority, outcome, notes, is_false_alarm, is_validated, is_closed"]],
  widths=[1.3, 1.8, 3.4])
add_para("")
add_para("Enumerations keep data consistent: categories (crime, fire, accident, disease_outbreak, cyclone_flood, "
         "drought, other), statuses (pending, verified, in_progress, resolved, fake), and authority types "
         "(police, fire_department, health, civil_protection).")

# ----------------------------------------------------------------------
# 6. TECH STACK
# ----------------------------------------------------------------------
add_heading("6. Technology Stack (and why)", 1)
add_table(["Layer", "Technology", "Why I chose it"],
  [["Backend framework", "FastAPI (Python)", "Fast, modern, automatic API docs, built-in validation"],
   ["Validation", "Pydantic", "Rejects bad input automatically before it reaches the logic"],
   ["Database / ORM", "SQLite + SQLAlchemy", "Zero-config, file-based, perfect for a prototype; ORM keeps schema in code"],
   ["Real-time", "WebSockets", "Push live updates instantly instead of constant polling"],
   ["Authentication", "JWT (python-jose) + bcrypt (passlib)", "Industry-standard tokens and one-way password hashing"],
   ["Frontend", "React 19 + Vite", "Component-based, fast dev server, fast production build"],
   ["Maps", "Leaflet + React-Leaflet", "Free, open-source interactive maps (no API key)"],
   ["HTTP client", "Axios", "Clean API calls from the browser"],
   ["Deployment", "Docker + Compose + Nginx + Traefik", "Reproducible, containerised, auto-HTTPS, production-ready"]],
  widths=[1.5, 2.0, 3.0])

# ----------------------------------------------------------------------
# 7. INNOVATION
# ----------------------------------------------------------------------
add_heading("7. The Innovation: Intelligent Multi-Authority Routing", 1)
add_para("This is the feature to emphasise for the Innovation & Creativity (20 marks) section.")
add_para("A naive system maps one category → one department. Real emergencies are messier: “a head-on collision "
         "and the cars caught fire, two people are trapped and bleeding” is an accident, but it actually needs "
         "Police + Fire Brigade + Ambulance all at once.")
add_para("My routing engine (services/notification.py) does three things:")
add_number(" every category has a default set of responsible authorities (e.g. fire → Fire Brigade and Health).", "Category mapping —")
add_number(" it scans the title and description for dozens of keywords and phrases (“burning”, “explosion”, "
           "“bleeding”, “trapped”, “robbery”, “flooding”, “collapse”…) and adds any extra authority "
           "those words imply, even if the reporter chose the wrong category. Multi-word phrases are matched as substrings; "
           "single words use word-boundary regex so “fire” doesn’t wrongly match “firefighter”.", "Keyword intelligence —")
add_number(" an accident described with severe wording (fire, trapped, explosion) automatically escalates to bring "
           "all three of Police, Ambulance and Fire on scene.", "Severity escalation —")
add_para("Then services/authority_offices.py takes the incident’s GPS coordinates and uses the Haversine great-circle "
         "formula to find the nearest actual office (from a directory of real Zimbabwean police stations, fire brigades, "
         "hospitals and Civil Protection offices) for each responsible authority — so the dispatcher closest to the scene is paged.")
add_para("Education 5.0 link (the form asks for this):", bold=True, space_after=2)
add_bullet(" an original, locally-built emergency-tech solution tailored to Zimbabwe’s authorities and emergency numbers.", "Innovation & Industrialisation —")
add_bullet(" directly addresses a real national public-safety gap.", "Community outreach / problem-solving —")
add_bullet(" applies geospatial computation (Haversine) and keyword-based classification to a practical domain.", "Research —")

# ----------------------------------------------------------------------
# 8. ASSESSMENT MAPPING
# ----------------------------------------------------------------------
add_heading("8. Mapping to the Assessment Form (Defence Cheat-Sheet)", 1)

def section_block(title, marks, points, say=None, say_label="Say:"):
    add_heading(f"{title}  —  {marks} marks", 2, GREEN)
    for pt in points:
        add_bullet(pt)
    if say:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(say_label + " ")
        r.bold = True
        r2 = p.add_run(say)
        r2.italic = True
        r2.font.color.rgb = BLUE

section_block("① Technical Development & Implementation", 25,
  ["Full incident lifecycle: report → auto-route → manage → respond → resolve / flag fake.",
   "Clean layered architecture (routers / services / schemas / models).",
   "Real database with 5 related tables and proper relationships.",
   "Complete CRUD on incidents, alerts, reports, authorities.",
   "Smart routing algorithm + geospatial nearest-office computation.",
   "Real-time WebSocket updates across all clients."],
  "“The system implements every functional requirement end-to-end, organised into a clean three-tier architecture with separated concerns.”")

section_block("② User Interface, Usability & Input Validation", 10,
  ["Two purpose-built, intuitive React interfaces (public vs. authority) with maps, icons and clear navigation.",
   "Input validation on every layer: Pydantic enforces minimum lengths, latitude ∈ [-90, 90], longitude ∈ [-180, 180], "
   "allowed severity values via regex, allowed outcomes, password minimum length — invalid data is rejected with a clear error, never crashing the server.",
   "In-app toasts and live notifications guide the user.",
   "Auto-generated interactive API documentation at /docs (your ‘help and documentation’)."],
  "“Validation happens at the API boundary with Pydantic, so the server cannot be crashed by malformed input — it returns a clear 422 error instead.”")

section_block("③ Testing & Evaluation", 20,
  ["Error handling: every endpoint guards against missing records (404), bad input (422), and wrong credentials (401). "
   "The WebSocket and event bus are wrapped in try/except and ‘never raise’, so a failed broadcast can’t take down a request.",
   "Fault tolerance: the frontend WebSocket auto-reconnects with exponential backoff if the connection drops; the backend "
   "restart: unless-stopped policy auto-restarts a crashed container.",
   "Response time / scalability: FastAPI is asynchronous; queries are indexed on primary keys and incident_id; the frontend is served as pre-built static files via Nginx.",
   "Tested manually across all user journeys; the /docs interface allows live endpoint testing."],
  "“It uses SQLite, which is excellent for this prototype scale; for very high concurrent national load the next step would be PostgreSQL — the ORM means that’s a one-line change.”",
  say_label="Be honest about limits:")

section_block("④ System Security", 10,
  ["Authentication: username/password login issuing JWT tokens that expire after 24 hours.",
   "Passwords are never stored in plain text — they are hashed with bcrypt (one-way, salted).",
   "Authorization / data scoping: each authority only sees incidents relevant to its type.",
   "Input sanitisation via Pydantic; the ORM uses parameterised queries which prevent SQL injection.",
   "HTTPS everywhere with auto-renewing Let’s Encrypt certificates (encryption in transit); HTTP force-redirected to HTTPS.",
   "Backup/recovery: database and uploads live on persistent Docker volumes that survive restarts and redeploys."],
  "“For production I would move the JWT secret key into an environment variable, tighten CORS from the permissive demo "
  "setting to the known domains, and enforce auth middleware on the management endpoints — the structure is already in place for all three.”",
  say_label="Honest hardening note (shows maturity):")

section_block("⑤ Innovation & Creativity", 20,
  ["Intelligent multi-authority keyword + severity routing (Section 7) — beyond a simple lookup table.",
   "Geospatial nearest-office dispatch using the Haversine formula.",
   "Real-time collaborative dashboards via WebSockets.",
   "Localised for Zimbabwe (real emergency numbers, real cities and stations).",
   "Clear Education 5.0 alignment (innovation, industrialisation, community outreach, research)."])

section_block("⑥ Documentation & Report Quality", 10,
  ["README with setup instructions, this defence guide, auto-generated API docs, clean commit history, and a Chapter 4 generator script."])

section_block("⑦ Presentation & Demonstration", 5,
  ["See the live demo script in Section 10."])

# ----------------------------------------------------------------------
# 9. CHALLENGES
# ----------------------------------------------------------------------
add_heading("9. Challenges Faced (and How I Solved Them)", 1)
add_para("Panels love this question. Have these ready:")
challenges = [
 ("Routing to the wrong / only one department.",
  "A single category can’t capture a complex emergency. Solution: built the keyword + severity escalation engine that returns multiple authorities from the incident text."),
 ("Real-time updates without hammering the server.",
  "Constantly polling the API is wasteful and slow. Solution: implemented a WebSocket event bus so the server pushes changes once and all clients update instantly."),
 ("Sending events from synchronous code.",
  "The REST endpoints are synchronous but WebSocket broadcasting is asynchronous. Solution: a small event-bus indirection (events.py) plus run_coroutine_threadsafe schedules the broadcast safely onto the running event loop, with a decoupled publisher so there are no circular imports."),
 ("Evolving the database without losing data.",
  "I needed to add new columns (validation/close flags) after data already existed. Solution: lightweight automatic migrations that run on startup and add missing columns only if absent."),
 ("Finding the nearest responder, not just a responder.",
  "Distance on a globe isn’t straight-line. Solution: implemented the Haversine great-circle formula over a directory of real Zimbabwean offices."),
 ("Deploying three apps under HTTPS cleanly.",
  "Three apps, one server, certificates, routing. Solution: Docker Compose + Traefik with automatic Let’s Encrypt certificates and per-domain routing labels."),
 ("CORS and separate frontends.",
  "Browsers block cross-origin API calls. Solution: configured CORS middleware and injected the API/WebSocket URLs at build time via Vite build args, so the same code works locally and in production."),
 ("Frontend dependency conflicts on React 19.",
  "Some packages lagged behind React 19. Solution: installed with --legacy-peer-deps in the Docker build to resolve peer-dependency conflicts cleanly."),
]
for i, (prob, sol) in enumerate(challenges, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(prob + "  ")
    r.bold = True
    p.add_run(sol)

# ----------------------------------------------------------------------
# 10. DEMO SCRIPT
# ----------------------------------------------------------------------
add_heading("10. Live Demo Script (for the 5-mark Presentation)", 1)
add_para("Run this exact sequence — it shows the whole system in 3–4 minutes:")
demo = [
 "Open the public portal. Show the home page and emergency numbers directory.",
 "Report an incident — pick “Accident”, type “Head-on collision on Seke Road, a car caught fire and someone is bleeding”, drop a pin, attach a photo, submit.",
 "Switch to the authority dashboard (logged in as Police). The bell notification pops up instantly (this is the WebSocket) — and Fire and Health also received it because of the keywords “fire” and “bleeding”. This is the innovation moment.",
 "Open the dashboard map — the new incident appears live.",
 "Update its status to “In Progress” — show it change.",
 "Issue a public alert (“Avoid Seke Road — accident and fire”) and show it appear on the public View Alerts page.",
 "File a responder report marking the outcome, then validate and close it — show the incident move to Resolved.",
 "Open Reports/Analytics — show the aggregated counts by category/status/day.",
 "Mention the live URLs and that it’s fully deployed under HTTPS.",
]
for step in demo:
    add_number(step)

# ----------------------------------------------------------------------
# 11. LIKELY QUESTIONS
# ----------------------------------------------------------------------
add_heading("11. Likely Panel Questions — Quick Answers", 1)
qa = [
 ("Why SQLite and not MySQL/Postgres?", "Right tool for a prototype: zero-config, file-based, fast. The SQLAlchemy ORM means switching to PostgreSQL for national scale is a configuration change, not a rewrite."),
 ("How is it secure?", "JWT auth, bcrypt-hashed passwords, parameterised ORM queries (no SQL injection), HTTPS in transit, per-authority data scoping."),
 ("What happens on bad input?", "Pydantic rejects it at the boundary with a 422 — the server never crashes."),
 ("What if the server restarts?", "Data is on Docker volumes; the container auto-restarts; clients auto-reconnect."),
 ("How does it know which authority to alert?", "Category mapping + keyword detection + severity escalation, then nearest-office by Haversine distance."),
 ("Is it real-time?", "Yes — WebSocket broadcasting, no polling."),
 ("What would you improve next?", "Real SMS gateway (Twilio/Africa’s Talking — currently simulated), PostgreSQL, env-var secrets, tighter CORS, automated tests, and an AI image classifier to auto-verify photos (the hook for that already exists)."),
 ("What’s innovative about it?", "The multi-authority intelligent routing and geospatial nearest-office dispatch — most reporting systems just file a ticket; mine decides who responds and from where."),
]
for q, a in qa:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Q: " + q + "  ")
    r.bold = True
    r.font.color.rgb = NAVY
    r2 = p.add_run("A: " + a)
    r2.italic = True

# ----------------------------------------------------------------------
# 12. SUMMARY
# ----------------------------------------------------------------------
add_heading("12. One-Paragraph Summary (memorise this)", 1)
p = doc.add_paragraph()
r = p.add_run("“This is a centralised disaster management web platform for Zimbabwe with three parts — a public "
  "reporting portal, an authority command dashboard, and a FastAPI backend — connected in real time over WebSockets. "
  "The public report incidents with location and photos; an intelligent routing engine reads the report, uses keyword and "
  "severity analysis plus category mapping to alert every relevant authority (Police, Fire, Health, Civil Protection), and "
  "uses the Haversine formula to page the nearest office. Authorities manage incidents on a live map, issue public alerts, "
  "and close the loop with responder reports. It’s secured with JWT and bcrypt, validated with Pydantic, and deployed as "
  "Docker containers behind a Traefik reverse proxy with automatic HTTPS.”")
r.italic = True
r.font.size = Pt(11.5)
r.font.color.rgb = BLUE

out = r"C:\Users\mistafotisevheni\Downloads\feedback\Project Defence Guide.docx"
doc.save(out)
print("Saved:", out)
